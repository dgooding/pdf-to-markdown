from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import random
import shutil
import stat
import time
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DEFAULT_OUT_DIR = ROOT / "tests" / "fixtures" / "generated"


@dataclass
class FixtureSpec:
    fixture_id: str
    filename: str
    source_format: str
    group: str
    description: str


class SyntheticDataFactory:
    """Deterministic synthetic value generation with fixed seed support."""

    def __init__(self, seed: int) -> None:
        self.seed = seed
        self.rng = random.Random(seed)
        self.first_names = [
            "Ari",
            "Jordan",
            "Taylor",
            "Morgan",
            "Riley",
            "Casey",
            "Alex",
            "Parker",
            "Drew",
            "Skyler",
        ]
        self.last_names = [
            "Anders",
            "Blake",
            "Chen",
            "Diaz",
            "Ellis",
            "Foster",
            "Garcia",
            "Hayes",
            "Irwin",
            "James",
        ]
        self.departments = [
            "Platform Reliability",
            "Workflow Automation",
            "Knowledge Operations",
            "Service Performance",
            "Incident Quality",
        ]
        self.products = [
            "Atlas Portal",
            "Northwind Console",
            "Blue Harbor Service",
            "Nimbus Workflow",
        ]
        self.statuses = ["open", "in_progress", "resolved", "blocked", "review"]

    def pick(self, values: Sequence[str]) -> str:
        return values[self.rng.randrange(0, len(values))]

    def person_name(self) -> str:
        return f"{self.pick(self.first_names)} {self.pick(self.last_names)}"

    def department(self) -> str:
        return self.pick(self.departments)

    def product_name(self) -> str:
        return self.pick(self.products)

    def document_id(self, prefix: str = "DOC") -> str:
        return f"{prefix}-{self.rng.randint(1000, 9999)}-{self.rng.randint(10, 99)}"

    def ticket_id(self, prefix: str = "INC") -> str:
        return f"{prefix}-{self.rng.randint(100000, 999999)}"

    def change_id(self) -> str:
        return f"CHG-{self.rng.randint(100000, 999999)}"

    def invoice_id(self) -> str:
        return f"INV-{self.rng.randint(10000, 99999)}"

    def hostname(self) -> str:
        return f"host-{self.rng.randint(10, 999)}.example.com"

    def email(self) -> str:
        first = self.pick(self.first_names).lower()
        last = self.pick(self.last_names).lower()
        return f"{first}.{last}{self.rng.randint(1,99)}@example.com"

    def url(self, path: str = "docs") -> str:
        return f"https://example.com/{path}/{self.rng.randint(1,999)}"

    def ipv4_doc(self) -> str:
        return f"192.0.2.{self.rng.randint(1, 254)}"

    def ipv6_doc(self) -> str:
        return f"2001:db8::{self.rng.randint(1, 65535):x}"

    def amount(self) -> str:
        return f"${self.rng.randint(10, 9000)}.{self.rng.randint(0, 99):02d}"

    def percentage(self) -> str:
        return f"{self.rng.randint(0, 100)}%"

    def decimal(self) -> str:
        return f"{self.rng.randint(0, 999)}.{self.rng.randint(0, 99):02d}"

    def version(self) -> str:
        return f"v{self.rng.randint(1,9)}.{self.rng.randint(0,9)}.{self.rng.randint(0,9)}"

    def date(self) -> str:
        year = 2025 + self.rng.randint(0, 2)
        month = self.rng.randint(1, 12)
        day = self.rng.randint(1, 28)
        return f"{year:04d}-{month:02d}-{day:02d}"

    def time(self) -> str:
        hour = self.rng.randint(0, 23)
        minute = self.rng.randint(0, 59)
        return f"{hour:02d}:{minute:02d}"

    def timezone_name(self) -> str:
        return self.pick(["UTC", "US/Eastern", "Europe/London", "Asia/Tokyo"])

    def status(self) -> str:
        return self.pick(self.statuses)

    def priority(self) -> str:
        return self.pick(["P1", "P2", "P3", "P4"])

    def severity(self) -> str:
        return self.pick(["critical", "high", "medium", "low"])


class CorpusGenerator:
    def __init__(
        self,
        out_dir: Path,
        seed: int,
        groups: Optional[List[str]] = None,
        cleanup: bool = False,
    ) -> None:
        self.out_dir = out_dir
        self.seed = seed
        self.data = SyntheticDataFactory(seed)
        self.groups = groups or ["docx", "pdf"]
        self.cleanup = cleanup

        self.docs_dir = self.out_dir / "documents"
        self.assets_dir = self.out_dir / "assets"
        self.expected_dir = self.out_dir / "expected"
        self.previews_dir = self.out_dir / "previews"

        self.generated: List[Dict[str, Any]] = []

    def run(self) -> Path:
        if self.cleanup and self.out_dir.exists():
            self._safe_rmtree(self.out_dir)

        self.docs_dir.mkdir(parents=True, exist_ok=True)
        self.assets_dir.mkdir(parents=True, exist_ok=True)
        self.expected_dir.mkdir(parents=True, exist_ok=True)
        self.previews_dir.mkdir(parents=True, exist_ok=True)

        visual_assets = self._create_visual_assets()

        if "docx" in self.groups:
            self._generate_docx_fixtures(visual_assets)
        if "pdf" in self.groups:
            self._generate_pdf_fixtures(visual_assets)

        corpus_manifest = {
            "generated_at_utc": datetime.now(timezone.utc).isoformat(),
            "synthetic": True,
            "seed": self.seed,
            "groups": self.groups,
            "total_fixtures": len(self.generated),
            "fixtures": self.generated,
        }
        manifest_path = self.out_dir / "generated-corpus.json"
        manifest_path.write_text(json.dumps(corpus_manifest, indent=2), encoding="utf-8")
        return manifest_path

    def _safe_rmtree(self, target: Path) -> None:
        def _onerror(func, path, _exc_info):
            try:
                os.chmod(path, stat.S_IWRITE)
                func(path)
            except Exception:
                pass

        # Retry for transient Windows file locks.
        last_error: Exception | None = None
        for _ in range(5):
            try:
                shutil.rmtree(target, onerror=_onerror)
                return
            except Exception as exc:  # noqa: BLE001
                last_error = exc
                time.sleep(0.25)
        if last_error is not None:
            raise last_error

    # ------------------------- shared helpers -------------------------
    def _file_sha256(self, path: Path) -> str:
        h = hashlib.sha256()
        h.update(path.read_bytes())
        return h.hexdigest()

    def _asset_record(self, name: str, path: Path, meaningful: bool, intended_caption: str) -> Dict[str, Any]:
        with Image.open(path) as img:
            width, height = img.size
            image_format = img.format
        return {
            "name": name,
            "path": f"assets/{path.name}",
            "format": image_format,
            "width": width,
            "height": height,
            "sha256": self._file_sha256(path),
            "meaningful": meaningful,
            "intended_caption": intended_caption,
        }

    def _safe_filename(self, raw: str) -> str:
        keep = []
        for ch in raw:
            if ch.isalnum() or ch in ("-", "_", ".", " ", "(", ")", "&"):
                keep.append(ch)
            else:
                keep.append("_")
        name = "".join(keep).strip()
        return name.replace("  ", " ")

    def _write_expected(self, fixture_id: str, payload: Dict[str, Any]) -> Path:
        expected_path = self.expected_dir / f"{fixture_id.lower()}.expected.json"
        expected_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
        return expected_path

    def _capture_preview(self, source: Path, fixture_id: str) -> Optional[str]:
        if source.suffix.lower() != ".pdf":
            return None
        preview_path = self.previews_dir / f"{fixture_id.lower()}-p1.png"
        doc = fitz.open(source)
        try:
            if doc.page_count > 0:
                pix = doc.load_page(0).get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
                pix.save(str(preview_path))
                return str(preview_path.relative_to(self.out_dir)).replace("\\", "/")
            return None
        finally:
            doc.close()

    def _register_fixture(
        self,
        spec: FixtureSpec,
        source_path: Path,
        expected_path: Path,
        expected_payload: Dict[str, Any],
    ) -> None:
        preview = self._capture_preview(source_path, spec.fixture_id)
        page_count = expected_payload.get("intended_page_count")
        entry = {
            "fixture_id": spec.fixture_id,
            "filename": source_path.name,
            "source_format": spec.source_format,
            "group": spec.group,
            "description": spec.description,
            "source_path": str(source_path.relative_to(self.out_dir)).replace("\\", "/"),
            "expected_sidecar": str(expected_path.relative_to(self.out_dir)).replace("\\", "/"),
            "sha256": self._file_sha256(source_path),
            "page_count": page_count,
            "synthetic": True,
        }
        if preview:
            entry["preview"] = preview
        self.generated.append(entry)

    def _docx_hyperlink(self, paragraph, text: str, url: str) -> None:
        part = paragraph.part
        rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel)

        run = OxmlElement("w:r")
        run_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        run_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_pr.append(underline)
        run.append(run_pr)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    # ------------------------- asset generation -------------------------
    def _create_visual_assets(self) -> Dict[str, Dict[str, Any]]:
        assets: Dict[str, Dict[str, Any]] = {}

        def save_image(name: str, draw_fn, meaningful: bool, caption: str) -> None:
            path = self.assets_dir / name
            draw_fn(path)
            assets[name] = self._asset_record(name, path, meaningful, caption)

        save_image("chart_bar.png", self._draw_bar_chart, True, "Synthetic bar chart")
        save_image("chart_line.png", self._draw_line_chart, True, "Synthetic line chart")
        save_image("chart_pie.png", self._draw_pie_chart, True, "Synthetic pie chart")
        save_image("diagram_flow.png", self._draw_flow_diagram, True, "Synthetic process flowchart")
        save_image("screenshot_mock.png", self._draw_screenshot_mock, True, "Synthetic application screenshot")
        save_image("photo_gradient.jpg", self._draw_photo_like, True, "Synthetic photograph-like gradient")
        save_image("logo_mark.png", self._draw_logo_mark, False, "Decorative geometric mark")
        save_image("status_cards.png", self._draw_status_cards, True, "Synthetic dashboard status cards")
        save_image("warning_icon.png", self._draw_warning_icon, False, "Decorative warning icon")
        save_image("info_icon.png", self._draw_info_icon, False, "Decorative information icon")

        # visually similar but binary-different
        path1 = self.assets_dir / "similar_a.png"
        path2 = self.assets_dir / "similar_b.png"
        self._draw_similar_image(path1, variation=1)
        self._draw_similar_image(path2, variation=2)
        assets[path1.name] = self._asset_record(path1.name, path1, True, "Similar synthetic image A")
        assets[path2.name] = self._asset_record(path2.name, path2, True, "Similar synthetic image B")

        return assets

    def _draw_canvas(self, path: Path, width: int = 1200, height: int = 700, bg=(255, 255, 255)) -> Tuple[Image.Image, ImageDraw.ImageDraw]:
        img = Image.new("RGB", (width, height), bg)
        draw = ImageDraw.Draw(img)
        return img, draw

    def _default_font(self):
        return ImageFont.load_default()

    def _draw_bar_chart(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 1200, 700, (247, 250, 255))
        f = self._default_font()
        draw.rounded_rectangle((25, 25, 1175, 675), radius=20, fill=(255, 255, 255), outline=(210, 220, 235), width=3)
        draw.text((45, 40), "Synthetic Bar Chart", fill=(20, 40, 80), font=f)
        left, top, right, bottom = 130, 140, 1110, 610
        draw.line((left, bottom, right, bottom), fill=(70, 90, 130), width=3)
        draw.line((left, top, left, bottom), fill=(70, 90, 130), width=3)
        values = [25, 38, 52, 47, 69, 58]
        x = left + 40
        for idx, val in enumerate(values):
            h = int((bottom - top) * val / 80)
            draw.rounded_rectangle((x, bottom - h, x + 80, bottom), radius=8, fill=(54, 123, 245))
            draw.text((x + 20, bottom + 10), f"M{idx+1}", fill=(30, 50, 80), font=f)
            draw.text((x + 20, bottom - h - 18), str(val), fill=(30, 50, 80), font=f)
            x += 135
        img.save(path)

    def _draw_line_chart(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 1200, 700, (251, 252, 255))
        f = self._default_font()
        draw.text((40, 36), "Synthetic Line Chart", fill=(15, 38, 70), font=f)
        left, top, right, bottom = 120, 130, 1100, 620
        draw.line((left, bottom, right, bottom), fill=(60, 90, 130), width=3)
        draw.line((left, top, left, bottom), fill=(60, 90, 130), width=3)
        points = []
        for i in range(12):
            px = left + i * (right - left) // 11
            py = bottom - int((math.sin(i / 2.2) * 0.35 + 0.45) * (bottom - top))
            points.append((px, py))
        draw.line(points, fill=(232, 104, 24), width=4)
        for px, py in points:
            draw.ellipse((px - 5, py - 5, px + 5, py + 5), fill=(232, 104, 24))
        img.save(path)

    def _draw_pie_chart(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 900, 700, (250, 252, 255))
        f = self._default_font()
        draw.text((30, 24), "Synthetic Pie Chart", fill=(21, 44, 88), font=f)
        box = (120, 120, 680, 680)
        slices = [120, 90, 75, 75]
        colors = [(54, 123, 245), (80, 180, 120), (250, 165, 52), (200, 88, 190)]
        start = 0
        for size, color in zip(slices, colors):
            draw.pieslice(box, start=start, end=start + size, fill=color, outline=(255, 255, 255), width=2)
            start += size
        draw.text((710, 180), "A 33%", fill=(20, 40, 80), font=f)
        draw.text((710, 220), "B 25%", fill=(20, 40, 80), font=f)
        draw.text((710, 260), "C 21%", fill=(20, 40, 80), font=f)
        draw.text((710, 300), "D 21%", fill=(20, 40, 80), font=f)
        img.save(path)

    def _draw_flow_diagram(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 1300, 720, (249, 251, 255))
        f = self._default_font()
        draw.text((30, 20), "Synthetic Flow Diagram", fill=(20, 45, 90), font=f)
        nodes = [
            (80, 120, 280, 200, "Start"),
            (400, 120, 650, 220, "Validate Input"),
            (760, 110, 1020, 220, "Decision"),
            (410, 330, 660, 430, "Transform"),
            (770, 330, 1030, 430, "Fallback"),
            (1080, 220, 1240, 320, "End"),
        ]
        for x1, y1, x2, y2, label in nodes:
            draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill=(255, 255, 255), outline=(70, 90, 140), width=3)
            draw.text((x1 + 24, y1 + 35), label, fill=(32, 52, 92), font=f)
        arrows = [
            (280, 160, 400, 170),
            (650, 170, 760, 165),
            (900, 220, 900, 330),
            (660, 380, 770, 380),
            (1030, 380, 1120, 280),
        ]
        for x1, y1, x2, y2 in arrows:
            draw.line((x1, y1, x2, y2), fill=(60, 90, 140), width=3)
            draw.polygon([(x2, y2), (x2 - 10, y2 - 6), (x2 - 10, y2 + 6)], fill=(60, 90, 140))
        img.save(path)

    def _draw_screenshot_mock(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 1400, 900, (20, 26, 36))
        f = self._default_font()
        draw.rectangle((0, 0, 1400, 68), fill=(28, 37, 55))
        draw.text((24, 24), "Synthetic Service Console", fill=(220, 230, 245), font=f)
        draw.rounded_rectangle((28, 90, 460, 860), radius=16, fill=(32, 43, 62), outline=(65, 84, 115), width=2)
        draw.text((50, 116), "Filters", fill=(170, 190, 230), font=f)
        for i in range(8):
            y = 160 + i * 80
            draw.rounded_rectangle((56, y, 430, y + 52), radius=8, fill=(22, 32, 50), outline=(66, 86, 122), width=1)
            draw.text((72, y + 18), f"Field {i+1}", fill=(150, 176, 220), font=f)
        draw.rounded_rectangle((500, 90, 1366, 500), radius=16, fill=(32, 43, 62), outline=(65, 84, 115), width=2)
        draw.text((520, 114), "Conversion Throughput", fill=(170, 190, 230), font=f)
        left, top, right, bottom = 540, 170, 1326, 470
        draw.line((left, bottom, right, bottom), fill=(95, 122, 170), width=3)
        draw.line((left, top, left, bottom), fill=(95, 122, 170), width=3)
        points = []
        for i in range(10):
            px = left + i * ((right - left) // 9)
            py = bottom - int((math.sin(i / 1.8) * 0.33 + 0.45) * (bottom - top))
            points.append((px, py))
        draw.line(points, fill=(105, 212, 255), width=4)
        draw.rounded_rectangle((500, 530, 1366, 860), radius=16, fill=(32, 43, 62), outline=(65, 84, 115), width=2)
        for i in range(7):
            draw.text((532, 570 + i * 38), f"#{8000+i} Synthetic activity message", fill=(190, 210, 245), font=f)
        img.save(path)

    def _draw_photo_like(self, path: Path) -> None:
        w, h = 1280, 720
        img = Image.new("RGB", (w, h), (0, 0, 0))
        draw = ImageDraw.Draw(img)
        for y in range(h):
            r = int(40 + (y / h) * 120)
            g = int(80 + (y / h) * 100)
            b = int(130 + (y / h) * 60)
            draw.line((0, y, w, y), fill=(r, g, b))
        draw.ellipse((200, 160, 560, 520), outline=(245, 245, 245), width=4)
        draw.rectangle((700, 180, 1050, 500), outline=(240, 240, 240), width=3)
        img.save(path, format="JPEG", quality=92)

    def _draw_logo_mark(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 512, 512, (255, 255, 255))
        draw.polygon([(90, 420), (260, 70), (430, 420)], fill=(48, 120, 230))
        draw.polygon([(160, 370), (260, 165), (360, 370)], fill=(255, 255, 255))
        img.save(path)

    def _draw_status_cards(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 1280, 720, (245, 248, 253))
        f = self._default_font()
        colors = [(87, 155, 252), (94, 186, 125), (245, 167, 59), (236, 102, 114)]
        labels = ["Open", "Resolved", "Pending", "Escalated"]
        for i, (c, label) in enumerate(zip(colors, labels)):
            x1 = 60 + i * 305
            x2 = x1 + 255
            draw.rounded_rectangle((x1, 120, x2, 300), radius=20, fill=c)
            draw.text((x1 + 24, 150), label, fill=(255, 255, 255), font=f)
            draw.text((x1 + 24, 220), str(100 + i * 33), fill=(255, 255, 255), font=f)
        img.save(path)

    def _draw_warning_icon(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 256, 256, (255, 255, 255))
        draw.polygon([(128, 24), (232, 220), (24, 220)], fill=(255, 193, 7))
        draw.rectangle((120, 95, 136, 165), fill=(70, 70, 70))
        draw.ellipse((118, 180, 138, 200), fill=(70, 70, 70))
        img.save(path)

    def _draw_info_icon(self, path: Path) -> None:
        img, draw = self._draw_canvas(path, 256, 256, (255, 255, 255))
        draw.ellipse((20, 20, 236, 236), fill=(33, 150, 243))
        draw.rectangle((122, 100, 134, 188), fill=(255, 255, 255))
        draw.ellipse((120, 64, 136, 80), fill=(255, 255, 255))
        img.save(path)

    def _draw_similar_image(self, path: Path, variation: int) -> None:
        img, draw = self._draw_canvas(path, 640, 360, (250, 252, 255))
        f = self._default_font()
        draw.rectangle((60, 60, 580, 300), outline=(60, 90, 140), width=4)
        draw.text((90, 90), "Synthetic Similar Visual", fill=(20, 45, 80), font=f)
        draw.line((90, 180, 540, 180 + variation), fill=(200, 80 + variation, 120), width=3)
        draw.ellipse((260, 140, 360, 240), outline=(40, 120, 210), width=3)
        draw.text((90, 250), f"variation={variation}", fill=(20, 45, 80), font=f)
        img.save(path)

    # ------------------------- DOCX fixtures -------------------------
    def _generate_docx_fixtures(self, assets: Dict[str, Dict[str, Any]]) -> None:
        builders = [
            self._docx_001,
            self._docx_002,
            self._docx_003,
            self._docx_004,
            self._docx_005,
            self._docx_006,
            self._docx_007,
            self._docx_008,
            self._docx_009,
            self._docx_010,
        ]
        for build in builders:
            spec, source, expected = build(assets)
            expected_path = self._write_expected(spec.fixture_id, expected)
            self._register_fixture(spec, source, expected_path, expected)

    def _docx_base_meta(self, fixture_id: str) -> Dict[str, Any]:
        return {
            "fixture_id": fixture_id,
            "synthetic": True,
            "seed": self.seed,
            "generated_data_note": "All content is synthetic and reproducible from seed.",
            "source_format": "docx",
            "expected_semantic_representation": "markdown",
        }

    def _docx_001(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec(
            fixture_id="DOCX-001",
            filename=self._safe_filename("DOCX-001 simple semantic document.docx"),
            source_format="docx",
            group="docx",
            description="Simple semantic document with heading hierarchy, lists, links, and metadata-like content.",
        )
        path = self.docs_dir / spec.filename
        doc = Document()

        doc.core_properties.author = self.data.person_name()
        doc.core_properties.title = "Synthetic Semantic Reference"

        title = doc.add_heading("Synthetic Document Title", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle = doc.add_paragraph("Synthetic subtitle for semantic conversion validation")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_heading("Section 1", level=2)
        p = doc.add_paragraph("Synthetic author: ")
        p.add_run(doc.core_properties.author).bold = True
        doc.add_paragraph(f"Created: {self.data.date()} Revised: {self.data.date()}")

        doc.add_heading("Section 1.1", level=3)
        p2 = doc.add_paragraph()
        p2.add_run("Bold").bold = True
        p2.add_run(" + ")
        p2.add_run("Italic").italic = True
        p2.add_run(" + ")
        bi = p2.add_run("BoldItalic")
        bi.bold = True
        bi.italic = True
        u = p2.add_run(" + Underlined")
        u.underline = True

        doc.add_paragraph("Simple bullet item", style="List Bullet")
        doc.add_paragraph("Simple numbered item", style="List Number")

        linkp = doc.add_paragraph("Reference: ")
        self._docx_hyperlink(linkp, "Example Link", "https://example.com/docs/semantic")

        doc.add_page_break()
        doc.add_heading("Section 2", level=2)
        doc.add_paragraph("Page break continuity paragraph.")

        section = doc.sections[0]
        section.header.paragraphs[0].text = "Synthetic Header"
        section.footer.paragraphs[0].text = "Synthetic Footer"

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Synthetic Document Title"],
                "expected_heading_hierarchy": ["#", "##", "###"],
                "expected_paragraphs": [
                    "Synthetic subtitle for semantic conversion validation",
                    "Page break continuity paragraph.",
                ],
                "expected_list_items": ["Simple bullet item", "Simple numbered item"],
                "expected_links": [{"text": "Example Link", "url": "https://example.com/docs/semantic"}],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": ["Synthetic Header"],
                "expected_page_footers": ["Synthetic Footer"],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": ["Bold", doc.core_properties.author],
                "expected_italic_spans": ["Italic"],
                "expected_underlined_spans": ["Underlined"],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["Synthetic Header", "Synthetic Footer"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _docx_002(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-002", self._safe_filename("DOCX-002 advanced lists.docx"), "docx", "docx", "Advanced nested list patterns.")
        path = self.docs_dir / spec.filename
        doc = Document()

        doc.add_heading("Advanced Lists", level=1)
        doc.add_paragraph("List preface paragraph.")
        doc.add_paragraph("Top bullet A", style="List Bullet")
        doc.add_paragraph("Nested bullet A.1", style="List Bullet 2")
        doc.add_paragraph("Nested bullet A.2", style="List Bullet 2")
        doc.add_paragraph("Top bullet B", style="List Bullet")

        doc.add_paragraph("Number item 1", style="List Number")
        doc.add_paragraph("Number item 1.a", style="List Number 2")
        doc.add_paragraph("Number item 2", style="List Number")

        p = doc.add_paragraph("Wrapped list text continuation style sample that should remain attached to the same logical list item when converted to markdown.", style="List Bullet")
        p.runs[0].bold = True

        mix = doc.add_paragraph("Mixed list with link", style="List Number")
        self._docx_hyperlink(mix, " example.com", "https://example.com/lists")

        doc.add_paragraph("A normal paragraph between list blocks.")
        doc.add_paragraph("Continuation list item after paragraph", style="List Number")

        # Custom-looking bullet glyph content
        doc.add_paragraph("▪ Symbol glyph bullet content")
        doc.add_paragraph("○ Hollow circle glyph content")

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Advanced Lists"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["List preface paragraph.", "A normal paragraph between list blocks."],
                "expected_list_items": [
                    "Top bullet A",
                    "Nested bullet A.1",
                    "Number item 1",
                    "Continuation list item after paragraph",
                ],
                "expected_links": [{"text": "example.com", "url": "https://example.com/lists"}],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": ["Wrapped list text continuation style sample"],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": ["Custom bullet glyphs may normalize differently by engine."],
            }
        )
        return spec, path, expected

    def _docx_003(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-003", self._safe_filename("DOCX-003 advanced tables.docx"), "docx", "docx", "Advanced tables with mixed cell content and merges.")
        path = self.docs_dir / spec.filename
        doc = Document()
        doc.add_heading("Advanced Tables", level=1)
        doc.add_paragraph("Paragraph above table should remain paragraph text.")

        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Date"
        hdr[2].text = "Amount"
        hdr[3].text = "Percent"
        hdr[4].text = "Notes"

        rows = [
            ("Alpha", self.data.date(), self.data.amount(), self.data.percentage(), "Line 1\nLine 2"),
            ("Beta", self.data.date(), self.data.amount(), self.data.percentage(), ""),
            ("Gamma", self.data.date(), self.data.amount(), self.data.percentage(), "Contains hyperlink"),
        ]
        for item, datev, amount, pct, notes in rows:
            row = table.add_row().cells
            row[0].text = item
            row[1].text = datev
            row[2].text = amount
            row[3].text = pct
            row[4].text = notes

        # Merge horizontal and vertical example
        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(2, 4).merge(table.cell(3, 4))

        cp = table.cell(3, 2).paragraphs[0]
        cp.add_run("BoldCell").bold = True
        cp.add_run("/")
        cp.add_run("ItalicCell").italic = True

        doc.add_paragraph("Paragraph below table should remain paragraph text.")
        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Advanced Tables"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": [
                    "Paragraph above table should remain paragraph text.",
                    "Paragraph below table should remain paragraph text.",
                ],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [
                    {"kind": "complex", "columns": 5, "header": ["Item", "Date", "Amount", "Percent", "Notes"]}
                ],
                "expected_table_cells": ["Alpha", "Beta", "Gamma", "BoldCell", "ItalicCell"],
                "expected_images": [],
                "expected_captions": ["Advanced Tables"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": ["BoldCell"],
                "expected_italic_spans": ["ItalicCell"],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_table_or_html_if_merge_complex",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Merged cells may require HTML table fallback."],
            }
        )
        return spec, path, expected

    def _docx_004(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-004", self._safe_filename("DOCX-004 images and captions.docx"), "docx", "docx", "Image placement and caption preservation with repeated assets.")
        path = self.docs_dir / spec.filename
        doc = Document()
        doc.add_heading("Images and Captions", level=1)

        doc.add_paragraph("Inline chart image:")
        doc.add_picture(str(self.assets_dir / "chart_bar.png"), width=Inches(5.5))
        doc.add_paragraph("Figure 1: Synthetic bar chart caption")

        doc.add_paragraph("Repeated image use on another section:")
        doc.add_picture(str(self.assets_dir / "chart_bar.png"), width=Inches(4.8))

        doc.add_paragraph("Transparent logo-like image:")
        doc.add_picture(str(self.assets_dir / "logo_mark.png"), width=Inches(2.0))

        doc.add_paragraph("Visually similar but binary-different pair:")
        doc.add_picture(str(self.assets_dir / "similar_a.png"), width=Inches(3.0))
        doc.add_picture(str(self.assets_dir / "similar_b.png"), width=Inches(3.0))

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 3,
                "expected_titles": ["Images and Captions"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Inline chart image:", "Repeated image use on another section:"],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [
                    "chart_bar.png",
                    "chart_bar.png",
                    "logo_mark.png",
                    "similar_a.png",
                    "similar_b.png",
                ],
                "expected_image_placements": [
                    {"asset": "chart_bar.png", "count": 2},
                    {"asset": "logo_mark.png", "count": 1},
                    {"asset": "similar_a.png", "count": 1},
                    {"asset": "similar_b.png", "count": 1},
                ],
                "expected_captions": ["Figure 1: Synthetic bar chart caption"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["chart_bar.png"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": ["Floating/inline placement may vary with DOCX renderer."],
            }
        )
        return spec, path, expected

    def _docx_005(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-005", self._safe_filename("DOCX-005 charts and business data.docx"), "docx", "docx", "Chart imagery + source data table.")
        path = self.docs_dir / spec.filename
        doc = Document()
        doc.add_heading("Charts and Business Data", level=1)

        doc.add_paragraph("Synthetic bar chart:")
        doc.add_picture(str(self.assets_dir / "chart_bar.png"), width=Inches(5.8))

        doc.add_paragraph("Synthetic line chart:")
        doc.add_picture(str(self.assets_dir / "chart_line.png"), width=Inches(5.8))

        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = "Series"
        table.rows[0].cells[1].text = "Jan"
        table.rows[0].cells[2].text = "Feb"
        table.rows[0].cells[3].text = "Mar"
        for name in ["North", "South", "Central"]:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = str(self.data.rng.randint(-30, 120))
            row[2].text = str(self.data.rng.randint(-30, 120))
            row[3].text = str(self.data.rng.randint(-30, 120))

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 3,
                "expected_titles": ["Charts and Business Data"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Synthetic bar chart:", "Synthetic line chart:"],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "source_data", "columns": 4}],
                "expected_images": ["chart_bar.png", "chart_line.png"],
                "expected_charts": ["bar", "line"],
                "expected_chart_labels": ["Jan", "Feb", "Mar", "North", "South", "Central"],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_chart_visual",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Charts are image-based in this fixture."],
            }
        )
        return spec, path, expected

    def _docx_006(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-006", self._safe_filename("DOCX-006 complex page layout.docx"), "docx", "docx", "Multi-section and mixed layout approximation fixture.")
        path = self.docs_dir / spec.filename
        doc = Document()

        doc.add_heading("Complex Page Layout", level=1)
        doc.add_paragraph("Portrait section paragraph.")
        doc.add_paragraph("Decorative separator simulation:")
        doc.add_paragraph("----------------------------------------")
        doc.add_paragraph("Callout: Validate that decorative separators do not trigger chart fallbacks.")

        doc.add_page_break()
        doc.add_heading("Section with alternate footer", level=2)
        section = doc.sections[-1]
        section.footer.paragraphs[0].text = "Layout Fixture Footer"
        section.header.paragraphs[0].text = "Layout Fixture Header"

        doc.add_paragraph("Pseudo two-column content line 1 | pseudo right column line 1")
        doc.add_paragraph("Pseudo two-column content line 2 | pseudo right column line 2")

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Complex Page Layout"],
                "expected_heading_hierarchy": ["#", "##"],
                "expected_paragraphs": [
                    "Portrait section paragraph.",
                    "Callout: Validate that decorative separators do not trigger chart fallbacks.",
                ],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": ["Layout Fixture Header"],
                "expected_page_footers": ["Layout Fixture Footer"],
                "expected_columns": [1, 2],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": ["decorative separators do not trigger chart fallbacks"],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["header", "footer"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["True DOCX multi-column fidelity depends on renderer."],
            }
        )
        return spec, path, expected

    def _docx_007(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-007", self._safe_filename("DOCX-007 code and technical content.docx"), "docx", "docx", "Code snippets and technical symbols fixture.")
        path = self.docs_dir / spec.filename
        doc = Document()
        doc.add_heading("Code and Technical Content", level=1)

        blocks = [
            "PowerShell: Get-ChildItem | Where-Object {$_.Length -gt 1024}",
            "Python: for i in range(3): print(i)",
            'JSON: {"service":"synthetic","ok":true}',
            "YAML: retries: 3\nmode: strict",
            "XML: <item id=\"1\">value</item>",
            "IPv4: 192.0.2.44 IPv6: 2001:db8::10",
            r"Windows Path: C:\\Synthetic\\App\\logs\\today.txt",
            "POSIX Path: /opt/synthetic/app/logs/today.txt",
            "Symbols: `pipe|asterisk*underscore_<>`",
        ]
        for b in blocks:
            p = doc.add_paragraph(b)
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(10)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Code and Technical Content"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": blocks,
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": blocks[:5],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": ["Inline escaping may differ by markdown generator."],
            }
        )
        doc.save(path)
        return spec, path, expected

    def _docx_008(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-008", self._safe_filename("DOCX-008 accessibility and structure.docx"), "docx", "docx", "Accessibility structure with deliberate heading jump.")
        path = self.docs_dir / spec.filename
        doc = Document()

        doc.add_heading("Accessibility and Structure", level=1)
        doc.add_heading("Primary Section", level=2)
        doc.add_paragraph("Descriptive link:")
        p = doc.add_paragraph()
        self._docx_hyperlink(p, "Open Synthetic Knowledge Base", "https://example.com/kb/1")
        doc.add_paragraph("Generic link:")
        p2 = doc.add_paragraph()
        self._docx_hyperlink(p2, "Click here", "https://example.com/generic")

        # Intentional jump
        doc.add_heading("Deliberate Jump Heading", level=4)

        doc.add_paragraph("Image with meaningful caption below:")
        doc.add_picture(str(self.assets_dir / "status_cards.png"), width=Inches(5.6))
        doc.add_paragraph("Figure: Status cards by queue state")

        doc.add_paragraph("Image without caption below:")
        doc.add_picture(str(self.assets_dir / "info_icon.png"), width=Inches(1.2))

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Accessibility and Structure"],
                "expected_heading_hierarchy": ["#", "##", "####"],
                "expected_paragraphs": ["Descriptive link:", "Generic link:"],
                "expected_list_items": [],
                "expected_links": [
                    {"text": "Open Synthetic Knowledge Base", "url": "https://example.com/kb/1"},
                    {"text": "Click here", "url": "https://example.com/generic"},
                ],
                "expected_tables": [],
                "expected_images": ["status_cards.png", "info_icon.png"],
                "expected_captions": ["Figure: Status cards by queue state"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Intentional heading-level jump should trigger warning."],
            }
        )
        return spec, path, expected

    def _docx_009(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-009", self._safe_filename("DOCX-009 unicode and special characters.docx"), "docx", "docx", "Unicode and symbol stability fixture.")
        path = self.docs_dir / spec.filename
        doc = Document()
        doc.add_heading("Unicode & Special Characters", level=1)

        lines = [
            "Accented: café naïve façade coöperate",
            "Quotes: “curly quotes” and apostrophe’s test",
            "Dashes: en–dash em—dash ellipsis…",
            "Math/Currency: ∑ π √2 € £ ¥ ₩",
            "Superscript/subscript approximation: x^2 and H2O",
            "Trademark-like: ™ ® ©",
            "Emoji: ✅ 📄 🔍",
            "Combining: A\u0301 e\u0301 i\u0301",
            "CJK sample: 変換テスト 漢字",
        ]
        for ln in lines:
            doc.add_paragraph(ln)

        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Unicode & Special Characters"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": lines,
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Bidirectional text not included by default in this environment."],
            }
        )
        return spec, path, expected

    def _docx_010(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("DOCX-010", self._safe_filename("DOCX-010 large synthetic document.docx"), "docx", "docx", "Large deterministic document for scale behavior.")
        path = self.docs_dir / spec.filename
        doc = Document()

        doc.add_heading("Large Synthetic Document", level=1)
        total_sections = 55
        repeated_image = self.assets_dir / "logo_mark.png"
        for i in range(1, total_sections + 1):
            doc.add_heading(f"Section {i}", level=2)
            doc.add_paragraph(
                f"Section {i} narrative for {self.data.department()} product {self.data.product_name()} status={self.data.status()} id={self.data.document_id('SEC')}."
            )
            doc.add_paragraph(f"List item A section {i}", style="List Bullet")
            doc.add_paragraph(f"List item B section {i}", style="List Bullet")
            if i % 5 == 0:
                t = doc.add_table(rows=2, cols=3)
                t.cell(0, 0).text = "Metric"
                t.cell(0, 1).text = "Value"
                t.cell(0, 2).text = "Status"
                t.cell(1, 0).text = "Queue"
                t.cell(1, 1).text = str(self.data.rng.randint(1, 999))
                t.cell(1, 2).text = self.data.status()
            if i % 7 == 0:
                doc.add_picture(str(repeated_image), width=Inches(1.0))
            if i < total_sections:
                doc.add_page_break()

        section = doc.sections[0]
        section.header.paragraphs[0].text = "Large Doc Header"
        section.footer.paragraphs[0].text = "Large Doc Footer"
        doc.save(path)

        expected = self._docx_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": total_sections,
                "expected_titles": ["Large Synthetic Document"],
                "expected_heading_hierarchy": ["#", "##"],
                "expected_paragraphs": ["Large Doc Header", "Large Doc Footer"],
                "expected_list_items": ["List item A section 1", "List item B section 1"],
                "expected_links": [],
                "expected_tables": [{"kind": "repeating", "every_n_sections": 5}],
                "expected_images": ["logo_mark.png"],
                "expected_captions": [],
                "expected_page_headers": ["Large Doc Header"],
                "expected_page_footers": ["Large Doc Footer"],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["header", "footer", "logo_mark.png"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Page count may vary slightly by renderer pagination."],
            }
        )
        return spec, path, expected

    # ------------------------- PDF fixtures -------------------------
    def _generate_pdf_fixtures(self, assets: Dict[str, Dict[str, Any]]) -> None:
        builders = [
            self._pdf_001,
            self._pdf_002,
            self._pdf_003,
            self._pdf_004,
            self._pdf_005,
            self._pdf_006,
            self._pdf_007,
            self._pdf_008,
            self._pdf_009,
            self._pdf_010,
            self._pdf_011,
            self._pdf_012,
            self._pdf_013,
            self._pdf_014,
            self._pdf_015,
            self._pdf_016,
            self._pdf_017,
            self._pdf_018,
            self._pdf_019,
            self._pdf_020,
        ]
        for build in builders:
            spec, source, expected = build(assets)
            expected_path = self._write_expected(spec.fixture_id, expected)
            self._register_fixture(spec, source, expected_path, expected)

    def _pdf_base_meta(self, fixture_id: str) -> Dict[str, Any]:
        return {
            "fixture_id": fixture_id,
            "synthetic": True,
            "seed": self.seed,
            "generated_data_note": "All content is synthetic and reproducible from seed.",
            "source_format": "pdf",
            "expected_semantic_representation": "markdown_with_targeted_fallback_when_unsafe",
        }

    def _new_pdf(self) -> fitz.Document:
        return fitz.open()

    def _pdf_add_header_footer(self, page: fitz.Page, header: str, footer: str) -> None:
        page.insert_text((40, 24), header, fontsize=9, fontname="helv", color=(0.4, 0.4, 0.4))
        page.insert_text((40, page.rect.height - 20), footer, fontsize=9, fontname="helv", color=(0.4, 0.4, 0.4))

    def _pdf_save(self, doc: fitz.Document, path: Path) -> None:
        doc.save(path)
        doc.close()

    def _pdf_001(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-001", self._safe_filename("PDF-001 native text.pdf"), "pdf", "pdf", "Native text, headings, lists, links, header/footer.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        self._pdf_add_header_footer(p, "Synthetic Header Native", "Synthetic Footer Native")
        p.insert_text((50, 70), "Native Text PDF", fontsize=24, fontname="helv")
        p.insert_text((50, 105), "Section One", fontsize=16, fontname="helv")
        p.insert_textbox(fitz.Rect(50, 130, 560, 210), "This is a native paragraph for semantic extraction validation.", fontsize=11)
        p.insert_text((70, 230), "• bullet one", fontsize=11)
        p.insert_text((70, 248), "• bullet two", fontsize=11)
        p.insert_text((70, 266), "1. numbered one", fontsize=11)
        link_rect = fitz.Rect(50, 300, 320, 316)
        p.insert_text((50, 312), "Example native hyperlink", fontsize=11, color=(0, 0, 1))
        p.insert_link({"kind": fitz.LINK_URI, "from": link_rect, "uri": "https://example.com/native"})
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Native Text PDF"],
                "expected_heading_hierarchy": ["#", "##"],
                "expected_paragraphs": ["This is a native paragraph for semantic extraction validation."],
                "expected_list_items": ["bullet one", "bullet two", "numbered one"],
                "expected_links": [{"text": "Example native hyperlink", "url": "https://example.com/native"}],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": ["Synthetic Header Native"],
                "expected_page_footers": ["Synthetic Footer Native"],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["blue", "black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["Synthetic Header Native", "Synthetic Footer Native"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_002(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-002", self._safe_filename("PDF-002 custom bullets.pdf"), "pdf", "pdf", "Custom bullet and nested list PDF.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        p.insert_text((50, 60), "Custom Bullets", fontsize=20)
        lines = [
            "• Standard bullet item",
            "▪ Square bullet item",
            "○ Hollow bullet item",
            "1. Numbered item",
            "   a. Lettered nested item",
            "   i. Roman nested item",
            "2. Wrapped list item with long content that should continue as same logical list entry after wrapping.",
        ]
        y = 100
        for ln in lines:
            p.insert_textbox(fitz.Rect(70, y, 560, y + 26), ln, fontsize=11)
            y += 24
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Custom Bullets"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": [],
                "expected_list_items": [
                    "Standard bullet item",
                    "Square bullet item",
                    "Hollow bullet item",
                    "Numbered item",
                ],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": ["Symbol bullets may normalize to standard markdown bullet."],
            }
        )
        return spec, path, expected

    def _pdf_003(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-003", self._safe_filename("PDF-003 ruled table.pdf"), "pdf", "pdf", "Ruled table with caption and surrounding paragraphs.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        p.insert_text((50, 60), "Ruled Table Fixture", fontsize=20)
        p.insert_text((50, 90), "Paragraph above ruled table.", fontsize=11)
        x0, y0 = 50, 120
        col_w = [150, 100, 100, 100]
        row_h = 30
        rows = 6
        cols = 4
        x = x0
        for w in col_w:
            p.draw_line(fitz.Point(x, y0), fitz.Point(x, y0 + row_h * rows), color=(0, 0, 0), width=1)
            x += w
        p.draw_line(fitz.Point(x, y0), fitz.Point(x, y0 + row_h * rows), color=(0, 0, 0), width=1)
        for r in range(rows + 1):
            y = y0 + r * row_h
            p.draw_line(fitz.Point(x0, y), fitz.Point(x0 + sum(col_w), y), color=(0, 0, 0), width=1)

        headers = ["Metric", "Date", "Amount", "Percent"]
        for i, h in enumerate(headers):
            p.insert_text((x0 + 10 + sum(col_w[:i]), y0 + 20), h, fontsize=10)
        for r in range(1, rows):
            p.insert_text((x0 + 10, y0 + r * row_h + 20), f"Row {r}", fontsize=10)
            p.insert_text((x0 + 160, y0 + r * row_h + 20), self.data.date(), fontsize=10)
            p.insert_text((x0 + 260, y0 + r * row_h + 20), self.data.amount(), fontsize=10)
            p.insert_text((x0 + 360, y0 + r * row_h + 20), self.data.percentage(), fontsize=10)

        p.insert_text((50, y0 + rows * row_h + 24), "Table 1: Synthetic ruled table", fontsize=10)
        p.insert_text((50, y0 + rows * row_h + 52), "Paragraph below ruled table.", fontsize=11)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Ruled Table Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Paragraph above ruled table.", "Paragraph below ruled table."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "ruled", "rows": 6, "columns": 4}],
                "expected_table_cells": ["Metric", "Date", "Amount", "Percent", "Row 1"],
                "expected_images": [],
                "expected_captions": ["Table 1: Synthetic ruled table"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "semantic_table_or_targeted_table_image",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_004(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-004", self._safe_filename("PDF-004 borderless table.pdf"), "pdf", "pdf", "Borderless coordinate-aligned table.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        p.insert_text((50, 60), "Borderless Table Fixture", fontsize=20)
        p.insert_text((50, 90), "Paragraph above borderless table.", fontsize=11)

        base_y = 130
        cols_x = [60, 220, 340, 460]
        headers = ["Name", "Qty", "Cost", "Owner"]
        for x, h in zip(cols_x, headers):
            p.insert_text((x, base_y), h, fontsize=11, fontname="helv")

        for r in range(1, 7):
            y = base_y + 22 * r
            p.insert_text((cols_x[0], y), f"Item-{r}", fontsize=10.5)
            p.insert_text((cols_x[1], y), str(self.data.rng.randint(1, 99)), fontsize=10.5)
            p.insert_text((cols_x[2], y), self.data.amount(), fontsize=10.5)
            p.insert_text((cols_x[3], y), self.data.person_name(), fontsize=10.5)

        p.insert_text((50, base_y + 190), "Paragraph below borderless table.", fontsize=11)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Borderless Table Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Paragraph above borderless table.", "Paragraph below borderless table."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "borderless", "rows": 7, "columns": 4}],
                "expected_table_cells": ["Name", "Qty", "Cost", "Owner", "Item-1"],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "semantic_or_targeted_table",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Borderless table detection confidence may vary."],
            }
        )
        return spec, path, expected

    def _pdf_005(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-005", self._safe_filename("PDF-005 complex table.pdf"), "pdf", "pdf", "Complex table with merged-cell simulation.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=792, height=612)
        p.insert_text((50, 50), "Complex Table Fixture", fontsize=22)
        p.insert_text((50, 80), "Grouped headers and merge-like layout simulation.", fontsize=11)

        # Simulated grouped header table
        x0, y0 = 50, 120
        width = 680
        height = 360
        p.draw_rect(fitz.Rect(x0, y0, x0 + width, y0 + height), color=(0, 0, 0), width=1.2)
        # horizontal lines
        for y in [160, 210, 260, 310, 360, 410, 460]:
            p.draw_line(fitz.Point(x0, y), fitz.Point(x0 + width, y), color=(0, 0, 0), width=1)
        # vertical lines (with merge gaps)
        for x in [180, 320, 470, 620]:
            p.draw_line(fitz.Point(x, y0), fitz.Point(x, y0 + height), color=(0, 0, 0), width=1)

        p.insert_text((60, 145), "Group A", fontsize=11)
        p.insert_text((330, 145), "Group B", fontsize=11)
        p.insert_text((60, 195), "Item", fontsize=10)
        p.insert_text((190, 195), "Value", fontsize=10)
        p.insert_text((330, 195), "Status", fontsize=10)
        p.insert_text((480, 195), "Notes", fontsize=10)

        y = 245
        for i in range(1, 6):
            p.insert_text((60, y), f"Row {i}", fontsize=10)
            p.insert_text((190, y), self.data.amount(), fontsize=10)
            p.insert_text((330, y), self.data.status(), fontsize=10)
            p.insert_text((480, y), f"Multiline note {i}", fontsize=10)
            y += 50

        p.insert_text((50, 510), "Footnote: merged-cell complexity expected.", fontsize=9)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Complex Table Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Grouped headers and merge-like layout simulation."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "complex", "requires_html_or_fallback": True}],
                "expected_table_cells": ["Group A", "Group B", "Row 1", "Status", "Notes"],
                "expected_images": [],
                "expected_captions": ["Footnote: merged-cell complexity expected."],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "html_or_targeted_table",
                "expected_review_requirement": "yes",
                "known_intentional_ambiguity": ["Merged-cell semantics may not map to markdown table."],
            }
        )
        return spec, path, expected

    def _pdf_006(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-006", self._safe_filename("PDF-006 vector chart.pdf"), "pdf", "pdf", "Vector-like chart region with paragraphs around it.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        p.insert_text((50, 60), "Vector Chart Fixture", fontsize=20)
        p.insert_text((50, 90), "Paragraph above chart region should remain semantic.", fontsize=11)

        # Draw chart-like geometry
        rect = fitz.Rect(70, 130, 540, 430)
        p.draw_rect(rect, color=(0.3, 0.4, 0.6), width=1)
        p.draw_line((100, 390), (500, 390), color=(0, 0, 0), width=2)
        p.draw_line((100, 160), (100, 390), color=(0, 0, 0), width=2)
        bars = [120, 170, 210, 150, 250]
        x = 130
        for b in bars:
            p.draw_rect(fitz.Rect(x, 390 - b, x + 45, 390), color=(0.2, 0.45, 0.9), fill=(0.4, 0.6, 0.95), width=1)
            p.insert_text((x + 5, 398), str(self.data.rng.randint(1, 9)), fontsize=8)
            x += 70
        p.insert_text((120, 145), "Synthetic chart title", fontsize=10)

        p.insert_textbox(
            fitz.Rect(50, 450, 560, 720),
            "Paragraph below chart region should remain semantic text and should not absorb chart labels or tick values.",
            fontsize=11,
        )
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Vector Chart Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": [
                    "Paragraph above chart region should remain semantic.",
                    "Paragraph below chart region should remain semantic text and should not absorb chart labels or tick values.",
                ],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_charts": ["bar"],
                "expected_chart_labels": ["Synthetic chart title"],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_chart_crop",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_007(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-007", self._safe_filename("PDF-007 multiple visuals.pdf"), "pdf", "pdf", "Multiple visual regions with body text separators.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        p = doc.new_page(width=792, height=612)
        p.insert_text((40, 40), "Multiple Visual Regions", fontsize=20)
        p.insert_text((40, 70), "Body paragraph before visuals.", fontsize=11)

        p.insert_image(fitz.Rect(40, 100, 360, 290), filename=str(self.assets_dir / "chart_bar.png"))
        p.insert_text((40, 300), "Figure 1: chart region", fontsize=10)

        p.insert_image(fitz.Rect(390, 100, 760, 290), filename=str(self.assets_dir / "chart_line.png"))
        p.insert_text((390, 300), "Figure 2: chart region", fontsize=10)

        # Decorative separator
        p.draw_line((40, 330), (760, 330), color=(0.5, 0.5, 0.5), width=1)

        # table-like
        p.insert_text((40, 360), "Metric   Q1   Q2", fontsize=10, fontname="cour")
        p.insert_text((40, 376), "Alpha    20   33", fontsize=10, fontname="cour")
        p.insert_text((40, 392), "Beta     18   29", fontsize=10, fontname="cour")

        p.insert_image(fitz.Rect(390, 345, 760, 560), filename=str(self.assets_dir / "photo_gradient.jpg"))
        p.insert_text((390, 572), "Figure 3: photo region", fontsize=10)

        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Multiple Visual Regions"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Body paragraph before visuals."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "text_grid", "rows": 3}],
                "expected_images": ["chart_bar.png", "chart_line.png", "photo_gradient.jpg"],
                "expected_captions": ["Figure 1: chart region", "Figure 2: chart region", "Figure 3: photo region"],
                "expected_charts": ["chart", "chart"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [2],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_multi_region",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_008(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-008", self._safe_filename("PDF-008 decorative vectors.pdf"), "pdf", "pdf", "Decorative vectors only, no chart.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=612, height=792)
        p.insert_text((50, 60), "Decorative Vectors", fontsize=20)
        p.insert_text((50, 92), "No chart should be inferred from decorative lines.", fontsize=11)

        p.draw_rect(fitz.Rect(40, 40, 572, 752), color=(0.6, 0.6, 0.6), width=1)  # border
        for y in [130, 220, 310, 400, 490, 580, 670]:
            p.draw_line((50, y), (560, y), color=(0.7, 0.7, 0.7), width=1)
        p.draw_rect(fitz.Rect(70, 150, 540, 210), color=(0.6, 0.6, 0.6), width=1, fill=(0.95, 0.95, 0.95))
        p.insert_text((80, 185), "Decorative text box", fontsize=11)

        self._pdf_save(doc, path)
        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Decorative Vectors"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["No chart should be inferred from decorative lines.", "Decorative text box"],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "no",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_009(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-009", self._safe_filename("PDF-009 two column.pdf"), "pdf", "pdf", "Two-column reading-order fixture.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=792, height=612)
        p.insert_text((60, 50), "Two-Column Fixture", fontsize=22)
        self._pdf_add_header_footer(p, "TwoCol Header", "TwoCol Footer")

        left = fitz.Rect(60, 90, 370, 550)
        right = fitz.Rect(420, 90, 730, 550)
        text_left = "\n".join([f"L{i}: synthetic left column sentence." for i in range(1, 20)])
        text_right = "\n".join([f"R{i}: synthetic right column sentence." for i in range(1, 20)])
        p.insert_textbox(left, text_left, fontsize=10.5)
        p.insert_textbox(right, text_right, fontsize=10.5)

        p.insert_image(fitz.Rect(430, 360, 720, 520), filename=str(self.assets_dir / "logo_mark.png"))
        p.insert_text((430, 535), "Figure: Column image caption", fontsize=10)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Two-Column Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["L1: synthetic left column sentence.", "R1: synthetic right column sentence."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["logo_mark.png"],
                "expected_captions": ["Figure: Column image caption"],
                "expected_page_headers": ["TwoCol Header"],
                "expected_page_footers": ["TwoCol Footer"],
                "expected_page_numbers": [],
                "expected_columns": [2],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["TwoCol Header", "TwoCol Footer"],
                "expected_fallback_representation": "targeted_for_figure_only",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Reading order must remain column-aware."],
            }
        )
        return spec, path, expected

    def _pdf_010(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-010", self._safe_filename("PDF-010 mixed orientation.pdf"), "pdf", "pdf", "Portrait, landscape, rotated page fixture.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        p1 = doc.new_page(width=612, height=792)
        p1.insert_text((50, 60), "Portrait Page", fontsize=22)
        p1.insert_text((50, 100), "Portrait narrative paragraph.", fontsize=11)

        p2 = doc.new_page(width=792, height=612)
        p2.insert_text((50, 60), "Landscape Page", fontsize=22)
        p2.insert_image(fitz.Rect(50, 120, 420, 420), filename=str(self.assets_dir / "status_cards.png"))
        p2.insert_text((50, 440), "Landscape table-ish text", fontsize=11)

        p3 = doc.new_page(width=612, height=792)
        p3.set_rotation(90)
        p3.insert_text((60, 60), "Rotated Page", fontsize=22)
        p3.insert_image(fitz.Rect(80, 120, 520, 500), filename=str(self.assets_dir / "photo_gradient.jpg"))
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 3,
                "expected_titles": ["Portrait Page", "Landscape Page", "Rotated Page"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Portrait narrative paragraph.", "Landscape table-ish text"],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "landscape_text_grid"}],
                "expected_images": ["status_cards.png", "photo_gradient.jpg"],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0, 0, 90],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_visual",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Rotation-normalized reading order may vary."],
            }
        )
        return spec, path, expected

    def _pdf_011(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-011", self._safe_filename("PDF-011 image-only scan.pdf"), "pdf", "pdf", "Image-only scan-like pages.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        # build scan-like image page from text drawn on raster
        scan_img = self.assets_dir / "scan_like_page.png"
        img, draw = self._draw_canvas(scan_img, 1240, 1754, (248, 248, 248))
        f = self._default_font()
        draw.text((120, 120), "Synthetic Scan Heading", fill=(20, 20, 20), font=f)
        for i in range(20):
            draw.text((120, 180 + i * 45), f"Scan paragraph line {i+1} with ticket {self.data.ticket_id()}.", fill=(30, 30, 30), font=f)
        draw.text((120, 1180), "• Scan bullet one", fill=(20, 20, 20), font=f)
        draw.text((120, 1225), "• Scan bullet two", fill=(20, 20, 20), font=f)
        img.save(scan_img)

        p = doc.new_page(width=612, height=792)
        p.insert_image(fitz.Rect(0, 0, 612, 792), filename=str(scan_img))
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Synthetic Scan Heading"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Scan paragraph line 1"],
                "expected_list_items": ["Scan bullet one", "Scan bullet two"],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["scan_like_page.png"],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": ["full_page"],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "visual_plus_ocr_if_available",
                "expected_review_requirement": "yes",
                "known_intentional_ambiguity": ["OCR availability is environment-dependent."],
            }
        )
        return spec, path, expected

    def _pdf_012(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-012", self._safe_filename("PDF-012 mixed native scanned.pdf"), "pdf", "pdf", "Mixed native + scanned pages.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        p1 = doc.new_page(width=612, height=792)
        p1.insert_text((50, 60), "Native page", fontsize=20)
        p1.insert_text((50, 95), "Native text should stay semantic.", fontsize=11)

        scan_img = self.assets_dir / "scan_like_page.png"
        p2 = doc.new_page(width=612, height=792)
        p2.insert_image(fitz.Rect(40, 60, 572, 740), filename=str(scan_img))

        p3 = doc.new_page(width=612, height=792)
        p3.insert_text((50, 60), "Native page with rasterized table", fontsize=16)
        p3.insert_image(fitz.Rect(50, 120, 560, 450), filename=str(self.assets_dir / "status_cards.png"))

        p4 = doc.new_page(width=612, height=792)
        p4.insert_text((50, 60), "Native with screenshot", fontsize=16)
        p4.insert_image(fitz.Rect(50, 120, 560, 550), filename=str(self.assets_dir / "screenshot_mock.png"))

        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 4,
                "expected_titles": ["Native page", "Native page with rasterized table", "Native with screenshot"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Native text should stay semantic."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "rasterized"}],
                "expected_images": ["scan_like_page.png", "status_cards.png", "screenshot_mock.png"],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": ["page_2", "raster_regions"],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "mixed_region_strategy",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Native/OCR dedup expected where overlaps exist."],
            }
        )
        return spec, path, expected

    def _pdf_013(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-013", self._safe_filename("PDF-013 image-only illustration.pdf"), "pdf", "pdf", "Image-only illustration page with optional caption.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        p = doc.new_page(width=612, height=792)
        p.insert_image(fitz.Rect(30, 80, 582, 700), filename=str(self.assets_dir / "diagram_flow.png"))
        p.insert_text((40, 728), "Figure: Synthetic process illustration", fontsize=10)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": [],
                "expected_heading_hierarchy": [],
                "expected_paragraphs": [],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["diagram_flow.png"],
                "expected_captions": ["Figure: Synthetic process illustration"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "visual_only",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": [],
            }
        )
        return spec, path, expected

    def _pdf_014(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-014", self._safe_filename("PDF-014 screenshot procedure.pdf"), "pdf", "pdf", "Synthetic screenshot procedure guide.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=792, height=612)
        p.insert_text((40, 40), "Screenshot Procedure", fontsize=20)
        p.insert_text((40, 68), "Step-by-step synthetic UI process", fontsize=11)

        p.insert_image(fitz.Rect(40, 90, 380, 330), filename=str(self.assets_dir / "screenshot_mock.png"))
        p.insert_textbox(fitz.Rect(410, 100, 760, 300), "1) Open synthetic console\n2) Select queue\n3) Click Convert\n4) Validate output", fontsize=11)
        p.insert_image(fitz.Rect(40, 350, 380, 580), filename=str(self.assets_dir / "status_cards.png"))
        p.insert_text((410, 380), "Callout A", fontsize=10)
        p.insert_text((410, 410), "Callout B", fontsize=10)

        self._pdf_save(doc, path)
        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Screenshot Procedure"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Step-by-step synthetic UI process"],
                "expected_list_items": ["Open synthetic console", "Select queue", "Click Convert", "Validate output"],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["screenshot_mock.png", "status_cards.png"],
                "expected_captions": ["Callout A", "Callout B"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [2],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": ["Callout A", "Callout B"],
                "expected_ocr_regions": ["screenshot_regions_if_text_missing"],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "targeted_screenshot_regions",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Callout order may remain ambiguous by layout extraction."],
            }
        )
        return spec, path, expected

    def _pdf_015(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-015", self._safe_filename("PDF-015 diagram flowchart.pdf"), "pdf", "pdf", "Diagram/flowchart fixture with support text.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p = doc.new_page(width=792, height=612)
        p.insert_text((40, 40), "Diagram and Flowchart", fontsize=20)
        p.insert_text((40, 72), "Supporting paragraph above diagram.", fontsize=11)
        p.insert_image(fitz.Rect(50, 100, 740, 500), filename=str(self.assets_dir / "diagram_flow.png"))
        p.insert_text((40, 540), "Supporting paragraph below diagram.", fontsize=11)
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 1,
                "expected_titles": ["Diagram and Flowchart"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Supporting paragraph above diagram.", "Supporting paragraph below diagram."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["diagram_flow.png"],
                "expected_captions": [],
                "expected_charts": ["diagram"],
                "expected_chart_labels": ["Start", "Validate Input", "Decision", "End"],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": ["diagram_labels_if_native_missing"],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "diagram_visual_with_labels_when_reliable",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Do not infer hidden node relationships."],
            }
        )
        return spec, path, expected

    def _pdf_016(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-016", self._safe_filename("PDF-016 links and annotations.pdf"), "pdf", "pdf", "Links and basic annotations fixture.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p1 = doc.new_page(width=612, height=792)
        p1.insert_text((50, 60), "Links and Annotations", fontsize=20)
        ext_rect = fitz.Rect(50, 100, 300, 116)
        p1.insert_text((50, 112), "External link text", fontsize=11, color=(0, 0, 1))
        p1.insert_link({"kind": fitz.LINK_URI, "from": ext_rect, "uri": "https://example.com/external"})

        int_rect = fitz.Rect(50, 140, 300, 156)
        p1.insert_text((50, 152), "Internal link text", fontsize=11, color=(0, 0, 1))

        p2 = doc.new_page(width=612, height=792)
        p2.insert_text((50, 60), "Target Page", fontsize=18)
        p2.insert_text((50, 90), "Internal link destination", fontsize=11)

        # Re-load page 0 after creating page 1 to avoid stale page-handle issues in PyMuPDF.
        p1_final = doc.load_page(0)
        p1_final.insert_link({"kind": fitz.LINK_GOTO, "from": int_rect, "page": 1, "to": fitz.Point(50, 60)})

        self._pdf_save(doc, path)
        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Links and Annotations", "Target Page"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Internal link destination"],
                "expected_list_items": [],
                "expected_links": [
                    {"text": "External link text", "url": "https://example.com/external"},
                    {"text": "Internal link text", "url": "#internal"},
                ],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["blue", "black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Annotation extraction support can vary by parser."],
            }
        )
        return spec, path, expected

    def _pdf_017(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-017", self._safe_filename("PDF-017 repeated asset.pdf"), "pdf", "pdf", "Repeated-asset placement fixture.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()

        p1 = doc.new_page(width=612, height=792)
        p1.insert_text((50, 60), "Repeated Asset Page 1", fontsize=18)
        p1.insert_image(fitz.Rect(80, 120, 300, 300), filename=str(self.assets_dir / "chart_pie.png"))

        p2 = doc.new_page(width=612, height=792)
        p2.insert_text((50, 60), "Distinct Asset Page 2", fontsize=18)
        p2.insert_image(fitz.Rect(80, 120, 300, 300), filename=str(self.assets_dir / "chart_line.png"))
        p2.insert_image(fitz.Rect(320, 120, 540, 300), filename=str(self.assets_dir / "similar_a.png"))

        p3 = doc.new_page(width=612, height=792)
        p3.insert_text((50, 60), "Repeated Asset Page 3", fontsize=18)
        p3.insert_image(fitz.Rect(80, 120, 300, 300), filename=str(self.assets_dir / "chart_pie.png"))
        p3.insert_image(fitz.Rect(320, 120, 540, 300), filename=str(self.assets_dir / "similar_b.png"))

        self._pdf_save(doc, path)
        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 3,
                "expected_titles": ["Repeated Asset Page 1", "Distinct Asset Page 2", "Repeated Asset Page 3"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": [],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": ["chart_pie.png", "chart_line.png", "similar_a.png", "similar_b.png"],
                "expected_image_placements": [
                    {"asset": "chart_pie.png", "pages": [1, 3]},
                    {"asset": "chart_line.png", "pages": [2]},
                ],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["chart_pie.png"],
                "expected_fallback_representation": "none",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Binary dedup allowed; placement must remain page-specific."],
            }
        )
        return spec, path, expected

    def _pdf_018(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-018", self._safe_filename("PDF-018 malformed and edge cases.pdf"), "pdf", "pdf", "Edge-case package (valid minimal fixture + blank page).")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        p1 = doc.new_page(width=612, height=792)
        p1.insert_text((50, 60), "Edge Case Fixture", fontsize=20)
        p2 = doc.new_page(width=612, height=792)  # blank page
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": 2,
                "expected_titles": ["Edge Case Fixture"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": [],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "review_or_blank_handling",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Blank page should not crash converter."],
            }
        )
        return spec, path, expected

    def _pdf_019(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-019", self._safe_filename("PDF-019 large synthetic.pdf"), "pdf", "pdf", "Large synthetic multi-page PDF.")
        path = self.docs_dir / spec.filename
        doc = self._new_pdf()
        total_pages = 100
        for i in range(1, total_pages + 1):
            p = doc.new_page(width=612, height=792)
            self._pdf_add_header_footer(p, "LargePDF Header", f"Page {i}")
            p.insert_text((50, 60), f"Large PDF Section {i}", fontsize=18)
            p.insert_textbox(
                fitz.Rect(50, 90, 560, 210),
                f"Synthetic page {i} for large benchmark. Ticket={self.data.ticket_id()} Change={self.data.change_id()} "
                f"Owner={self.data.person_name()} Department={self.data.department()}.",
                fontsize=10.5,
            )
            p.insert_text((70, 230), f"• item A {i}", fontsize=10)
            p.insert_text((70, 246), f"• item B {i}", fontsize=10)
            if i % 10 == 0:
                p.insert_image(fitz.Rect(60, 270, 300, 420), filename=str(self.assets_dir / "chart_line.png"))
            if i % 15 == 0:
                p.insert_image(fitz.Rect(320, 270, 560, 420), filename=str(self.assets_dir / "photo_gradient.jpg"))
            if i % 20 == 0:
                p.insert_text((50, 450), "Metric  Value", fontsize=10, fontname="cour")
                p.insert_text((50, 466), "alpha   41", fontsize=10, fontname="cour")
                p.insert_text((50, 482), "beta    52", fontsize=10, fontname="cour")
        self._pdf_save(doc, path)

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": total_pages,
                "expected_titles": ["Large PDF Section 1"],
                "expected_heading_hierarchy": ["#"],
                "expected_paragraphs": ["Synthetic page 1 for large benchmark."],
                "expected_list_items": ["item A 1", "item B 1"],
                "expected_links": [],
                "expected_tables": [{"kind": "periodic_text_grid", "every_n_pages": 20}],
                "expected_images": ["chart_line.png", "photo_gradient.jpg"],
                "expected_captions": [],
                "expected_page_headers": ["LargePDF Header"],
                "expected_page_footers": ["Page 1"],
                "expected_page_numbers": ["Page 1"],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": [],
                "expected_repeated_elements": ["LargePDF Header"],
                "expected_fallback_representation": "mixed_semantic_targeted",
                "expected_review_requirement": "possible",
                "known_intentional_ambiguity": ["Performance-focused fixture may vary in runtime by machine."],
            }
        )
        return spec, path, expected

    def _pdf_020(self, assets: Dict[str, Dict[str, Any]]) -> Tuple[FixtureSpec, Path, Dict[str, Any]]:
        spec = FixtureSpec("PDF-020", self._safe_filename("PDF-020 original regression sample.pdf"), "pdf", "pdf", "Alias fixture referencing original sample regression expectations.")
        path = self.docs_dir / spec.filename
        source_sample = ROOT / "file-sample_150kB.pdf"
        if source_sample.exists():
            shutil.copy2(source_sample, path)
            page_count = fitz.open(path).page_count
        else:
            # deterministic fallback if sample not present
            doc = self._new_pdf()
            p = doc.new_page(width=612, height=792)
            p.insert_text((50, 60), "Sample placeholder missing", fontsize=18)
            self._pdf_save(doc, path)
            page_count = 1

        expected = self._pdf_base_meta(spec.fixture_id)
        expected.update(
            {
                "intended_page_count": page_count,
                "expected_titles": ["Regression sample expectations"],
                "expected_heading_hierarchy": ["#", "##"],
                "expected_paragraphs": ["Page 1 semantic text with targeted chart expected."],
                "expected_list_items": [],
                "expected_links": [],
                "expected_tables": [{"kind": "targeted_table_fallback_expected"}],
                "expected_images": [],
                "expected_captions": [],
                "expected_page_headers": [],
                "expected_page_footers": [],
                "expected_page_numbers": [],
                "expected_columns": [1],
                "expected_rotations": [0],
                "expected_text_colors": ["black"],
                "expected_bold_spans": [],
                "expected_italic_spans": [],
                "expected_underlined_spans": [],
                "expected_code_blocks": [],
                "expected_callouts": [],
                "expected_ocr_regions": ["page_4_if_ocr_available"],
                "expected_repeated_elements": [],
                "expected_fallback_representation": "as_validated_prototype_v1",
                "expected_review_requirement": "yes",
                "known_intentional_ambiguity": ["Do not hard-code logic for this regression fixture."],
            }
        )
        return spec, path, expected


# ------------------------- validation helpers -------------------------
def validate_generated_corpus(corpus_manifest: Path) -> Dict[str, Any]:
    payload = json.loads(corpus_manifest.read_text(encoding="utf-8"))
    out_dir = corpus_manifest.parent
    fixtures = payload.get("fixtures", [])

    validation = {
        "manifest": str(corpus_manifest),
        "total_fixtures": len(fixtures),
        "valid_docx": 0,
        "valid_pdf": 0,
        "expected_sidecars_present": 0,
        "preview_images_present": 0,
        "filename_issues": [],
        "missing_files": [],
        "text_presence_checks": [],
        "ok": True,
    }

    for item in fixtures:
        source = out_dir / item["source_path"]
        expected = out_dir / item["expected_sidecar"]

        if not source.exists():
            validation["missing_files"].append(str(source))
            validation["ok"] = False
            continue

        if expected.exists():
            validation["expected_sidecars_present"] += 1
        else:
            validation["missing_files"].append(str(expected))
            validation["ok"] = False

        name = source.name
        if any(x in name for x in ["..", ":", "\\", "/"]):
            validation["filename_issues"].append(name)
            validation["ok"] = False

        suffix = source.suffix.lower()
        if suffix == ".docx":
            try:
                doc = Document(source)
                _ = doc.paragraphs[:3]
                validation["valid_docx"] += 1
                if expected.exists():
                    exp = json.loads(expected.read_text(encoding="utf-8"))
                    text_targets = exp.get("expected_titles", [])[:1] + exp.get("expected_paragraphs", [])[:1]
                    combined = "\n".join(p.text for p in doc.paragraphs)
                    for t in text_targets:
                        if t and t.split(" ")[0] in combined:
                            validation["text_presence_checks"].append({"fixture": item["fixture_id"], "token": t.split(" ")[0], "present": True})
                        else:
                            validation["text_presence_checks"].append({"fixture": item["fixture_id"], "token": t.split(" ")[0] if t else "", "present": False})
            except Exception:
                validation["ok"] = False
        elif suffix == ".pdf":
            try:
                pdf = fitz.open(source)
                page_count = pdf.page_count
                pdf.close()
                if page_count >= 1:
                    validation["valid_pdf"] += 1
                if expected.exists():
                    exp = json.loads(expected.read_text(encoding="utf-8"))
                    intended = exp.get("intended_page_count")
                    if isinstance(intended, int) and intended > 0 and page_count != intended:
                        # allow known pagination ambiguity notes for DOCX only; PDF should match exactly here
                        validation["ok"] = False
                        validation["text_presence_checks"].append({"fixture": item["fixture_id"], "page_count_expected": intended, "page_count_actual": page_count})
            except Exception:
                validation["ok"] = False

        preview = item.get("preview")
        if preview and (out_dir / preview).exists():
            validation["preview_images_present"] += 1

    return validation


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate deterministic synthetic advanced PDF/DOCX fixture corpus.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUT_DIR, help="Output directory for generated corpus.")
    parser.add_argument("--seed", type=int, default=20260730, help="Deterministic seed value.")
    parser.add_argument(
        "--groups",
        type=str,
        default="docx,pdf",
        help="Comma-separated groups to generate (docx,pdf).",
    )
    parser.add_argument("--cleanup", action="store_true", help="Delete output dir before generation.")
    parser.add_argument("--validate", action="store_true", help="Run fixture validation after generation.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    groups = [g.strip().lower() for g in args.groups.split(",") if g.strip()]
    valid_groups = {"docx", "pdf"}
    for g in groups:
        if g not in valid_groups:
            raise ValueError(f"Unsupported group: {g}")

    generator = CorpusGenerator(out_dir=args.output_dir, seed=args.seed, groups=groups, cleanup=args.cleanup)
    manifest_path = generator.run()

    print("Synthetic corpus generated")
    print(f"- output_dir: {args.output_dir}")
    print(f"- manifest: {manifest_path}")
    print(f"- total fixtures: {len(generator.generated)}")

    if args.validate:
        report = validate_generated_corpus(manifest_path)
        report_path = args.output_dir / "generation-validation.json"
        report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
        print(f"- validation: {report_path}")
        print(f"- validation ok: {report.get('ok')}")


if __name__ == "__main__":
    main()
