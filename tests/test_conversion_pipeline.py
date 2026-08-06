from __future__ import annotations

import json
import re
import tempfile
import unittest
from pathlib import Path
import shutil

import fitz
from docx import Document

from convert_to_md import ConversionContext, convert_file_to_markdown, write_markdown


class ConversionPipelineTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_root = Path(tempfile.mkdtemp(prefix="md-convert-tests-"))
        cls.fixtures = cls.temp_root / "fixtures"
        cls.fixtures.mkdir(parents=True, exist_ok=True)

        cls._make_native_pdf(cls.fixtures / "native.pdf")
        cls._make_heading_pdf(cls.fixtures / "headings.pdf")
        cls._make_bullets_pdf(cls.fixtures / "bullets.pdf")
        cls._make_image_only_pdf(cls.fixtures / "image_only.pdf")
        cls._make_tableish_pdf(cls.fixtures / "tableish.pdf")
        cls._make_docx(cls.fixtures / "sample.docx")
        (cls.fixtures / "sample txt file.txt").write_text("Hello\n- one\n- two\n", encoding="utf-8")

    @classmethod
    def _make_native_pdf(cls, path: Path) -> None:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 60), "Native PDF Test", fontsize=20)
        page.insert_text((50, 100), "This is a normal paragraph for extraction testing.", fontsize=11)
        doc.save(path)
        doc.close()

    @classmethod
    def _make_heading_pdf(cls, path: Path) -> None:
        doc = fitz.open()
        p = doc.new_page()
        p.insert_text((50, 70), "Main Title", fontsize=28)
        p.insert_text((50, 120), "Section Heading", fontsize=18)
        p.insert_text((50, 150), "Body text should not become heading.", fontsize=11)
        doc.save(path)
        doc.close()

    @classmethod
    def _make_bullets_pdf(cls, path: Path) -> None:
        doc = fitz.open()
        p = doc.new_page()
        p.insert_text((50, 70), "• First bullet", fontsize=11)
        p.insert_text((50, 90), "• Second bullet", fontsize=11)
        p.insert_text((50, 110), "• Third bullet", fontsize=11)
        doc.save(path)
        doc.close()

    @classmethod
    def _make_image_only_pdf(cls, path: Path) -> None:
        doc = fitz.open()
        p = doc.new_page()
        shape = p.new_shape()
        shape.draw_rect(fitz.Rect(80, 120, 520, 520))
        shape.finish(color=(0, 0, 0), fill=(0.2, 0.6, 0.9))
        shape.commit()
        doc.save(path)
        doc.close()

    @classmethod
    def _make_tableish_pdf(cls, path: Path) -> None:
        doc = fitz.open()
        p = doc.new_page()
        p.insert_text((50, 70), "Metric    Q1    Q2", fontsize=11, fontname="cour")
        p.insert_text((50, 90), "Open      10    12", fontsize=11, fontname="cour")
        p.insert_text((50, 110), "Closed    8     11", fontsize=11, fontname="cour")
        doc.save(path)
        doc.close()

    @classmethod
    def _make_docx(cls, path: Path) -> None:
        doc = Document()
        doc.add_heading("DOCX Fixture", level=1)
        p = doc.add_paragraph("This has a link: ")
        p.add_run("https://example.com")
        doc.add_paragraph("Item one", style="List Bullet")
        doc.add_paragraph("Item two", style="List Bullet")
        doc.save(path)

    def _convert(self, source: Path, mode: str = "hybrid") -> tuple[Path, Path]:
        out_dir = self.temp_root / f"out_{source.stem}_{mode}"
        assets = out_dir / "assets"
        assets.mkdir(parents=True, exist_ok=True)

        ctx = ConversionContext(
            output_dir=out_dir,
            assets_dir=assets,
            overwrite=True,
            pdf_mode=mode,
            tesseract_cmd=None,
            prefer_markitdown=False,
            strict_validation=False,
        )
        md = convert_file_to_markdown(source, ctx)
        md_path = write_markdown(out_dir, source, md, overwrite=True)
        return md_path, out_dir

    def test_pdf_basic_native_text(self) -> None:
        md_path, out_dir = self._convert(self.fixtures / "native.pdf", mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("# native", text.lower())
        self.assertIn("## Page 1", text)
        self.assertIn("Extracted Text", text)

        manifest = out_dir / "native-manifest.json"
        self.assertTrue(manifest.exists())
        payload = json.loads(manifest.read_text(encoding="utf-8"))
        self.assertTrue(payload["validation"]["passed"])

    def test_pdf_headings_inferred(self) -> None:
        md_path, _ = self._convert(self.fixtures / "headings.pdf", mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertRegex(text, r"###\s+Main Title|####\s+Main Title")

    def test_pdf_bullet_glyphs_normalized(self) -> None:
        md_path, _ = self._convert(self.fixtures / "bullets.pdf", mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("- First bullet", text)
        self.assertNotIn("", text)

    def test_pdf_image_only_fallback(self) -> None:
        md_path, out_dir = self._convert(self.fixtures / "image_only.pdf", mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("Figure from page 1", text)
        manifests = list(out_dir.glob("*-manifest.json"))
        self.assertTrue(manifests)
        manifest = json.loads(manifests[0].read_text(encoding="utf-8"))
        self.assertIn("regions_rendered", manifest)

    def test_pdf_visual_mode(self) -> None:
        md_path, _ = self._convert(self.fixtures / "native.pdf", mode="visual")
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("Visual-first mode enabled", text)

    def test_docx_conversion_keeps_structure(self) -> None:
        md_path, _ = self._convert(self.fixtures / "sample.docx")
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("# DOCX Fixture", text)
        self.assertIn("- Item one", text)

    def test_txt_conversion(self) -> None:
        src = self.fixtures / "sample txt file.txt"
        md_path, _ = self._convert(src)
        text = md_path.read_text(encoding="utf-8")
        self.assertIn("# sample txt file", text.lower())
        self.assertIn("- one", text)

    def test_special_filename_pdf(self) -> None:
        src = self.fixtures / "special name @123.pdf"
        shutil.copy2(self.fixtures / "native.pdf", src)
        try:
            md_path, out_dir = self._convert(src)
            self.assertTrue(md_path.exists())
            manifest = out_dir / "special-name-123-manifest.json"
            self.assertTrue(manifest.exists())
        finally:
            src.unlink(missing_ok=True)

    def test_integration_file_sample_if_present(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        md_path, out_dir = self._convert(workspace_sample, mode="hybrid")
        self.assertTrue(md_path.exists())
        manifests = list(out_dir.glob("*-manifest.json"))
        self.assertTrue(manifests)
        payload = json.loads(manifests[0].read_text(encoding="utf-8"))
        self.assertTrue(payload["validation"]["passed"])

    def test_vertical_slice_document_model_and_dual_status(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        _, out_dir = self._convert(workspace_sample, mode="hybrid")
        manifests = list(out_dir.glob("*-manifest.json"))
        self.assertTrue(manifests)
        payload = json.loads(manifests[0].read_text(encoding="utf-8"))

        self.assertIn("technical_status", payload)
        self.assertIn("fidelity_status", payload)
        self.assertIn(payload["technical_status"], {"passed", "failed"})
        self.assertIn(payload["fidelity_status"], {"high", "moderate", "low", "review_required"})

        doc_result = payload.get("document_result", {})
        self.assertIn("pages", doc_result)
        self.assertEqual(doc_result.get("page_count"), 4)
        self.assertTrue(all("candidates" in p for p in doc_result["pages"]))
        self.assertTrue(all("selected_candidate" in p for p in doc_result["pages"]))

    def test_vertical_slice_quality_report_exists(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        _, out_dir = self._convert(workspace_sample, mode="hybrid")
        reports = list(out_dir.glob("*-quality-report.json"))
        self.assertTrue(reports)
        report = json.loads(reports[0].read_text(encoding="utf-8"))
        self.assertIn("page_summaries", report)
        self.assertTrue(report["page_summaries"])
        self.assertTrue(all("candidate_scores" in p for p in report["page_summaries"]))

    def test_published_markdown_has_no_environment_diagnostics(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        md_path, _ = self._convert(workspace_sample, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertNotIn("OCR (Tesseract) is not available", text)
        self.assertNotIn("Text extraction quality was low", text)

    def test_targeted_fallback_preferred_over_full_page_when_table_present(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        _, out_dir = self._convert(workspace_sample, mode="hybrid")
        payload = json.loads(next(out_dir.glob("*-manifest.json")).read_text(encoding="utf-8"))
        page2 = [p for p in payload["document_result"]["pages"] if p["page_number"] == 2][0]
        rendered = payload.get("regions_rendered", [])
        self.assertNotIn("page-002-full", rendered)
        self.assertTrue(any(fr["fallback_type"] == "table_crop" for fr in page2.get("fallback_records", [])))

    def test_decorative_vectors_do_not_force_false_fullpage_fallback(self) -> None:
        md_path, out_dir = self._convert(self.fixtures / "headings.pdf", mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        payload = json.loads(next(out_dir.glob("*-manifest.json")).read_text(encoding="utf-8"))
        self.assertNotIn("full_page", str(payload.get("document_result", {})))
        self.assertNotIn("full-page", text.lower())

    def test_table_fallback_excludes_surrounding_paragraphs(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        md_path, _ = self._convert(workspace_sample, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        page2 = text.split("## Page 2", 1)[1].split("## Page 3", 1)[0]
        self.assertIn("page-002-table-001", page2)
        self.assertNotIn("page-002-full-page-001", page2)

    def test_full_page_and_embedded_duplication_detected_or_avoided(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        md_path, _ = self._convert(workspace_sample, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        page4 = text.split("## Page 4", 1)[1]
        # expect no duplicate full-page+embedded on page 4 after fix
        self.assertFalse("page-004-full-page-001" in page4 and "page-004-image-001" in page4)

    def test_bold_opening_not_deep_heading(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        md_path, _ = self._convert(workspace_sample, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        self.assertNotRegex(text, r"^######\s", msg="Should not produce deep heading level 6")

    def test_pdf_dedup_does_not_create_false_cross_page_image_placements(self) -> None:
        src = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not src.exists():
            self.skipTest("workspace sample PDF not present")

        md_path, out_dir = self._convert(src, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        manifests = list(out_dir.glob("*-manifest.json"))
        self.assertTrue(manifests, "Expected at least one manifest file in conversion output.")
        manifest_path = manifests[0]
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))

        placements = payload.get("image_placements", [])
        self.assertTrue(isinstance(placements, list))

        by_page: dict[int, set[str]] = {}
        for placement in placements:
            page = int(placement["source_page"])
            asset_path = placement["asset_path"]
            by_page.setdefault(page, set()).add(asset_path)

        page_refs: dict[int, set[str]] = {}
        current_page = None
        for line in text.splitlines():
            m = re.match(r"^##\s+Page\s+(\d+)", line)
            if m:
                current_page = int(m.group(1))
                page_refs.setdefault(current_page, set())
                continue
            if current_page is None:
                continue
            img = re.search(r"!\[[^\]]*\]\((assets/[^)]+)\)", line)
            if img:
                page_refs[current_page].add(img.group(1))

        for page, refs in page_refs.items():
            allowed = by_page.get(page, set())
            # page-wide fallback images are allowed independent of embedded-image placements
            fallback = {r for r in refs if "full-page" in r or "table-" in r}
            for ref in refs - fallback:
                self.assertIn(
                    ref,
                    allowed,
                    msg=f"Image reference {ref} appears under page {page} without a placement record on that page.",
                )

    def test_pdf_manifest_includes_ocr_provider_and_table_strategy_fields(self) -> None:
        workspace_sample = Path(__file__).resolve().parents[1] / "file-sample_150kB.pdf"
        if not workspace_sample.exists():
            self.skipTest("workspace sample PDF not present")
        _, out_dir = self._convert(workspace_sample, mode="hybrid")
        payload = json.loads(next(out_dir.glob("*-manifest.json")).read_text(encoding="utf-8"))

        self.assertIn("ocr_provider", payload)
        self.assertIn("available", payload["ocr_provider"])
        self.assertIn("tables_detected", payload)
        if payload["tables_detected"]:
            t0 = payload["tables_detected"][0]
            self.assertIn("strategy_selected", t0)
            self.assertIn("confidence", t0)
            self.assertIn("strategy_attempt_order", t0)

    def test_two_column_pdf_reading_order_prefers_column_grouping(self) -> None:
        src = self.fixtures / "two_col.pdf"
        doc = fitz.open()
        p = doc.new_page(width=792, height=612)
        p.insert_text((60, 60), "Two Column Test", fontsize=20)
        y = 110
        for i in range(1, 6):
            p.insert_text((70, y), f"L{i} left", fontsize=11)
            p.insert_text((430, y), f"R{i} right", fontsize=11)
            y += 24
        doc.save(src)
        doc.close()

        md_path, _ = self._convert(src, mode="hybrid")
        text = md_path.read_text(encoding="utf-8")
        # Ensure left column sequence is preserved before right-column sequence starts.
        self.assertTrue(text.find("L1 left") < text.find("L5 left"))
        self.assertTrue(text.find("L5 left") < text.find("R1 right"))



if __name__ == "__main__":
    unittest.main(verbosity=2)
